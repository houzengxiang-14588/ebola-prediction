"""将 项目总结报告.md 转换为 Word .doc 文件"""
import re, os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

os.chdir("D:/Ebola")

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.25

for i in range(1, 4):
    h_style = doc.styles[f"Heading {i}"]
    h_style.font.name = "Microsoft YaHei"
    h_style.font.color.rgb = RGBColor(0, 0, 0)

def add_para(text, bold=False, size=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    return p

def add_table_from_md(lines, start_idx):
    rows = []
    for line in lines[start_idx:]:
        line = line.strip()
        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if all(c.startswith("---") or c.startswith(":--") for c in cells):
                continue
            rows.append(cells)
        elif not line.startswith("|"):
            break

    if not rows:
        return start_idx + 1

    table = doc.add_table(rows=len(rows), cols=len(rows[0]), style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(cell_text)
            run.font.size = Pt(9)
            if i == 0:
                run.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    return start_idx + len(rows) + 1


def parse_inline(p, text):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            p.add_run(part)


def convert_md_to_doc(md_path):
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if line == "---" and i == 0:
            i += 1
            while i < len(lines) and lines[i].strip() != "---":
                i += 1
            i += 1
            continue

        if not line:
            doc.add_paragraph()
            i += 1
            continue

        if line.startswith("# "):
            add_para(line[2:], bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[3:], level=2)
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_heading(line[4:], level=3)
            i += 1
            continue

        if line.startswith("---"):
            doc.add_paragraph("─" * 60)
            i += 1
            continue

        if line.startswith("|"):
            i = add_table_from_md(lines, i)
            continue

        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i].rstrip())
                i += 1
            i += 1
            if code_lines:
                p = doc.add_paragraph()
                for cl in code_lines:
                    run = p.add_run(cl + "\n")
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue

        if line.startswith("- ") or line.startswith("* "):
            text = re.sub(r"^[-*] ", "", line)
            p = doc.add_paragraph(style="List Bullet")
            parse_inline(p, text)
            i += 1
            continue

        ol_match = re.match(r"^(\d+)\.\s+(.*)", line)
        if ol_match:
            text = ol_match.group(2)
            p = doc.add_paragraph(style="List Number")
            parts = re.split(r"(\*\*.*?\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            i += 1
            continue

        if line.startswith(">"):
            text = line[1:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        else:
            p = doc.add_paragraph()
            parse_inline(p, line)
        i += 1


if __name__ == "__main__":
    convert_md_to_doc("项目总结报告.md")
    doc.save("项目总结报告.doc")
    print("已生成: D:/Ebola/项目总结报告.doc")
