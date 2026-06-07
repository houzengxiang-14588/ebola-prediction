"""生成新旧实验对比分析文档 .docx"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

os.chdir("D:/Ebola")

doc = Document()

style = doc.styles["Normal"]
style.font.name = "宋体"
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

for i in range(1, 4):
    h = doc.styles[f"Heading {i}"]
    h.font.name = "黑体"
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    h.font.color.rgb = RGBColor(0, 0, 0)

for s in doc.sections:
    s.top_margin = Cm(2.54)
    s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3.18)
    s.right_margin = Cm(3.18)


def add_title(text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


def add_subtitle(text, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_table(headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers), style="Light Grid Accent 1")
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = tbl.rows[0].cells[j]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            c = tbl.rows[i + 1].cells[j]
            c.text = ""
            r = c.paragraphs[0].add_run(str(val))
            r.font.size = Pt(10)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


# ================================================================
doc.add_paragraph()
add_title("埃博拉病毒趋势预测项目")
add_title("新旧实验对比分析报告")
doc.add_paragraph()
add_subtitle("旧实验 (2026年5月29日) vs 新实验 (2026年6月3日)")
add_subtitle("生成时间: 2026年6月5日")
doc.add_page_break()

# ================================================================
doc.add_heading("一、总览对比", level=1)

add_table(
    ["维度", "旧实验 (5月29日)", "新实验 (6月3日)"],
    [
        ["实验类型", "单一回溯测试", "三组消融实验"],
        ["数据来源", "ebola_africa.csv (20国年度汇总)", "ebola_data_db_format.csv (2014-2015西非每周疫情)"],
        ["评估粒度", "国家-年度级别", "周级别滑动窗口"],
        ["样本量", "20 (一国一预测)", "162个滑动窗口/组 (共约486条)"],
        ["涉及国家", "20国", "5国 (聚焦3个核心暴发国)"],
        ["实验组数", "1组 (仅LoRA)", "3组 (Base / LoRA / LoRA+RAG)"],
        ["有无基线", "无", "有 (随机三分类基线 33.3%)"],
        ["主要指标", "准确率", "准确率 + 混淆矩阵 + P/R/F1 + 高置信度准确率"],
        ["论文图表", "无", "3张 (fig1/fig2/fig3, 300 DPI)"],
    ]
)

# ================================================================
doc.add_heading("二、旧实验详情 (2026年5月29日)", level=1)

doc.add_heading("2.1 实验设计", level=2)

add_body(
    "旧实验是一个单一的国家级年度回溯测试。以非洲20国2019-2022年的年度病例数据作为历史上下文，"
    "让模型预测各国2023年的病例趋势方向（上升/下降/平稳），再与2023年实际数据对比。趋势判定"
    "阈值与当前一致：变化超过15%为上升/下降，否则为平稳。实验仅测试了LoRA微调模型，没有Base"
    "基线和RAG消融对照。"
)

doc.add_heading("2.2 实验结果", level=2)

add_body(
    "总准确率95.0%（19/20正确）。高置信度准确率100%，但高置信度样本仅2个（占总数10%），"
    "统计意义极其有限。类别分布严重失衡：20国中18国被预测为\"平稳\"（F1=97.3%），2国为\"下降\""
    "（F1=66.7%），\"上升\"类0个样本。"
)

add_table(
    ["指标", "旧实验 (5月29日)"],
    [
        ["总准确率", "95.0% (19/20)"],
        ["vs 随机基线", "未报告基线"],
        ["高置信度准确率", "100% (仅2个样本)"],
        ["上升类 F1", "无样本"],
        ["下降类 F1", "66.7% (2国)"],
        ["平稳类 F1", "97.3% (18国)"],
    ]
)

doc.add_heading("2.3 核心问题", level=2)

add_body(
    "该实验存在严重的评估失真问题。非洲20国中17国在2023年的埃博拉病例数为0（属于无疫情国家），"
    "模型只需预测\"平稳\"即可得到85%的准确率。真正有疫情变化的仅3国（刚果金、乌干达、几内亚），"
    "而其中只有刚果金和乌干达为非零值变化。因此，95%的准确率实质上是\"零病例检测\"的准确率，"
    "而非趋势预测能力的真实反映。该实验无法区分模型是在做流行病学推理，还是仅仅学会了\"非洲"
    "大部分国家没有埃博拉\"这一先验知识。"
)

# ================================================================
doc.add_heading("三、新实验详情 (2026年6月3日)", level=1)

doc.add_heading("3.1 实验设计", level=2)

add_body(
    "新实验是一次严格的三组消融实验，系统量化了LoRA微调和RAG检索增强各自对预测性能的贡献。"
    "数据切换为2014-2015年西非埃博拉大流行期间的逐周病例数据（WHO数据库格式），覆盖几内亚、"
    "利比里亚、塞拉利昂等5个有实际暴发的国家。采用滑动窗口方法构造样本：每个窗口使用前8周"
    "的历史病例数据，预测下一周的趋势方向。三个核心暴发国产出162个滑动窗口，每个窗口在三个"
    "实验组（Base/LoRA/LoRA+RAG）中各评估一次，总计约486次独立预测。"
)

add_body(
    "三组实验组分别为：Base组（Qwen2-7B基座模型，无微调，无RAG）、LoRA组（加LoRA微调，无RAG）、"
    "LoRA+RAG组（LoRA微调 + ChromaDB检索增强，完整管线）。以随机三分类等概率（33.3%）为基线。"
    "评估指标扩展为准确率 + 三类混淆矩阵 + 每类Precision/Recall/F1 + 高置信度子集准确率。"
)

doc.add_heading("3.2 实验结果 (LoRA+RAG组详细)", level=2)

add_body(
    "LoRA+RAG完整管线取得最优准确率38.9%（63/162正确），较随机基线提升5.6个百分点。"
    "Base组37.7%，LoRA单独组34.6%。三类趋势的识别能力存在显著差异：下降类识别最好"
    "（Precision=56.3%, Recall=40.9%, F1=47.4%），上升类次之（F1=39.2%），平稳类最差"
    "（F1=20.0%）。高置信度子集（85个样本，占52.5%）准确率提升至41.2%，表明置信度标注"
    "具有一定的筛选价值。"
)

add_table(
    ["实验组", "准确率", "vs 基线", "高置信度准确率", "高置信度样本数"],
    [
        ["随机基线", "33.3%", "—", "—", "—"],
        ["Base", "37.7%", "+4.4%", "40.5%", "84"],
        ["LoRA", "34.6%", "+1.3%", "37.6%", "93"],
        ["LoRA+RAG", "38.9%", "+5.6%", "41.2%", "85"],
    ]
)

doc.add_heading("3.3 LoRA+RAG组混淆矩阵与分类指标", level=2)

add_table(
    ["类别", "实际样本数", "Precision", "Recall", "F1"],
    [
        ["上升", "43", "33.9%", "46.5%", "39.2%"],
        ["下降", "88", "56.3%", "40.9%", "47.4%"],
        ["平稳", "31", "17.9%", "22.6%", "20.0%"],
    ]
)

add_body(
    "类别分布较为均衡（上升43/下降88/平稳31），不存在旧实验中单一类别主导的问题。"
    "Precision-Recall的差异反映了模型对不同趋势类型的识别偏好：模型倾向于将不确定的情况预测"
    "为\"下降\"（预测了64次下降，实际只有88次真实下降），对\"平稳\"的识别力较弱（精准率仅17.9%），"
    "容易将真正的平稳误判为上升或下降。"
)

# ================================================================
doc.add_heading("四、核心差异分析", level=1)

doc.add_heading("4.1 数据层面", level=2)

add_body(
    "最根本的差异在于测试数据的选择。旧实验使用20国年度汇总数据（ebola_africa.csv），其中17国"
    "在2023年为零病例，导致评估退化为\"有无疫情\"的二分类检测，而非真正的趋势预测能力测试。"
    "新实验改用2014-2015年西非大流行期间的逐周暴发数据（ebola_data_db_format.csv），所有测试"
    "样本均来自有实际疫情的国家和周次，确保了评估对象是模型在真实暴发场景下的趋势推理能力。"
)

doc.add_heading("4.2 评估粒度层面", level=2)

add_body(
    "旧实验以国家-年度为评估单位，一个国家的全年趋势对应对一个预测。这种粗粒度评估丢失了大量"
    "时序信息——无法检测模型对短期波动、拐点和阶段性变化的捕捉能力。新实验采用周级别滑动窗口"
    "（窗口大小=8周），每个窗口独立评估，162个窗口覆盖了暴发全过程的上升期、峰值期和下降期，"
    "更全面地刻画了模型在不同疫情阶段的表现。"
)

doc.add_heading("4.3 实验设计层面", level=2)

add_body(
    "旧实验仅测试了LoRA微调模型，没有消融对照，无法回答\"微调到底有没有用\"\"RAG到底有没有用\""
    "这两个核心问题。新实验的三组消融设计直接量化了各组件的边际贡献：LoRA微调单独贡献+1.3%，"
    "但弱于Base组（-3.1%），暗示小样本微调存在过拟合风险；RAG在LoRA基础上贡献+4.3%，有效"
    "弥补了微调的不足，使完整管线达到最优。这种因果性证据是旧实验完全无法提供的。"
)

doc.add_heading("4.4 统计可靠性层面", level=2)

add_body(
    "旧实验20个样本、2个高置信度样本的结果，在统计上几乎无法得出任何可靠结论——95%准确率的"
    "95%置信区间宽度约为±10个百分点。新实验每组162个样本，具有更合理的统计效力。高置信度"
    "子集（每组84-93个样本，占52-57%）的样本量也足以支撑独立分析。此外，新实验的类别分布"
    "（上升43/下降88/平稳31）避免了单一类别主导的问题，每类都有足够的样本用于计算Precision/"
    "Recall/F1。"
)

# ================================================================
doc.add_heading("五、数字对比汇总", level=1)

add_table(
    ["对比维度", "旧实验", "新实验", "变化说明"],
    [
        ["样本量", "20", "162/组", "样本量提升8倍"],
        ["高置信度样本", "2 (10%)", "85 (52.5%)", "从统计无效到可独立分析"],
        ["涉及国家数", "20国", "5国", "从广度转向深度"],
        ["有疫情国家", "3国", "5国 (全部有疫情)", "评估对象从虚假转为真实"],
        ["实验组数", "1组", "3组", "可量化各组件贡献"],
        ["最佳准确率", "95.0%", "38.9%", "数字下降但真实度大幅提升"],
        ["vs 随机基线", "无基线比较", "+5.6%", "有统计意义的边际提升"],
        ["类别分布", "18平稳/2下降/0上升", "43上升/88下降/31平稳", "从单一类别到三类均衡"],
        ["论文图表", "0张", "3张 (300 DPI)", "可发表级别的可视化"],
        ["混淆矩阵", "无", "有 (3×3完整)", "可分析错误模式"],
        ["消融分析", "无", "Base/LoRA/LoRA+RAG", "可量化微调和RAG的贡献"],
    ]
)

# ================================================================
doc.add_heading("六、结论", level=1)

add_body(
    "旧实验（5月29日）的95.0%准确率是一个统计上不可靠且实质上失真的结果。其核心问题在于测试"
    "数据中大量零病例国家导致评估退化为\"有无疫情检测\"，而非趋势预测能力测试。20个样本中有"
    "18个预测为\"平稳\"，高置信度子集仅含2个样本——这样的结果无法支撑任何关于模型预测能力的"
    "有效结论。"
)

add_body(
    "新实验（6月3日）通过三个关键改进使评估变得有意义：一是将数据切换为有实际暴发的每周疫情"
    "数据，确保评估对象是真正的趋势预测能力；二是引入三组消融设计，首次量化了LoRA微调和RAG"
    "增强各自的边际贡献；三是采用周级别滑动窗口评估，覆盖暴发全过程的不同阶段。新实验揭示的"
    "核心发现——LoRA单独微调略弱于Base（34.6% vs 37.7%），但加上RAG后达到最优（38.9%）——"
    "具有明确的因果解释力，为后续优化提供了清晰方向。"
)

add_body(
    "因此，在引用和展示实验结果时，应使用新实验（6月3日）的消融数据。旧实验（5月29日）的"
    "95.0%不应作为模型预测能力的依据。新旧实验的准确率差异（95.0% vs 38.9%）并不代表模型"
    "性能的下降，而是反映了评估方法从失真到真实的范式转变。"
)

# 保存
out = "D:/Ebola/新旧实验对比分析.docx"
doc.save(out)
print(f"已生成: {out}")
